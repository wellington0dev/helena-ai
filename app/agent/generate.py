"""Tools de geração de mídia (CLAUDE.md §7): imagem, áudio (TTS), documento.

Cada handler gera o arquivo, salva no storage e insere uma `messages` do
assistant com a mídia — que o `runner` coleta ao final do turno. Retorna ao
modelo um resumo do que foi criado (não os bytes).
"""
import io
import wave

from flask import current_app
from google.genai import types

from app.agent.gemini import get_client
from app.extensions import db, write_lock
from app.media import storage
from app.models import Message

_TTS_RATE = 24000  # Gemini TTS: PCM L16 mono 24kHz
# Modelos de mídia (imagem/TTS) às vezes devolvem um candidato DEGENERADO
# (sem content/parts) — um no-op transitório. Reamostrar costuma resolver.
_MEDIA_ATTEMPTS = 3


def _new_assistant_media(user_id, content, media_url, media_type, media_meta) -> int:
    with write_lock:
        msg = Message(
            user_id=user_id,
            role="assistant",
            content=content or "",
            media_url=media_url,
            media_type=media_type,
            media_meta=media_meta,
        )
        db.session.add(msg)
        db.session.commit()
        return msg.id


def _pcm_to_wav(pcm: bytes, rate: int = _TTS_RATE) -> tuple[bytes, float]:
    """Embrulha PCM L16 mono em container WAV. Devolve (wav_bytes, duração_s)."""
    buf = io.BytesIO()
    with wave.open(buf, "wb") as w:
        w.setnchannels(1)
        w.setsampwidth(2)  # 16-bit
        w.setframerate(rate)
        w.writeframes(pcm)
    duration = round(len(pcm) / (rate * 2), 2)  # mono, 2 bytes/amostra
    return buf.getvalue(), duration


def _media_parts(client, model, contents, config) -> list | None:
    """Gera com um modelo de mídia e devolve as `parts` do candidato, retentando
    quando a resposta vem degenerada (sem candidato/content/parts — no-op
    transitório frequente em imagem/TTS). None se todas as tentativas falharem."""
    for _ in range(_MEDIA_ATTEMPTS):
        resp = client.models.generate_content(
            model=model, contents=contents, config=config
        )
        cand = resp.candidates[0] if resp.candidates else None
        parts = cand.content.parts if (cand and cand.content and cand.content.parts) else None
        if parts:
            return parts
    return None


# --------------------------------------------------------------------------- #
# Handlers
# --------------------------------------------------------------------------- #

def generate_image(user_id: int, args: dict) -> dict:
    prompt = (args.get("prompt") or "").strip()
    if not prompt:
        return {"ok": False, "error": "prompt vazio"}
    cfg = current_app.config
    client = get_client(cfg["GEMINI_API_KEY"])
    parts = _media_parts(
        client,
        cfg["GEMINI_IMAGE_MODEL"],
        prompt,
        types.GenerateContentConfig(response_modalities=["TEXT", "IMAGE"]),
    )
    if parts is None:
        return {"ok": False, "error": "geração bloqueada ou vazia"}

    img_bytes = None
    caption = ""
    for p in parts:
        if getattr(p, "inline_data", None) and p.inline_data.data:
            img_bytes = p.inline_data.data
        elif getattr(p, "text", None):
            caption = p.text.strip()
    if img_bytes is None:
        return {"ok": False, "error": "modelo não retornou imagem"}

    media_url = storage.save_bytes(user_id, img_bytes, "png")
    msg_id = _new_assistant_media(
        user_id, caption, media_url, "image",
        {"mime": "image/png", "prompt": prompt},
    )
    return {"ok": True, "message_id": msg_id, "created": "imagem"}


def generate_audio(user_id: int, args: dict) -> dict:
    text = (args.get("text") or "").strip()
    if not text:
        return {"ok": False, "error": "text vazio"}
    cfg = current_app.config
    client = get_client(cfg["GEMINI_API_KEY"])
    parts = _media_parts(
        client,
        cfg["GEMINI_TTS_MODEL"],
        text,
        types.GenerateContentConfig(
            response_modalities=["AUDIO"],
            speech_config=types.SpeechConfig(
                voice_config=types.VoiceConfig(
                    prebuilt_voice_config=types.PrebuiltVoiceConfig(
                        voice_name=cfg["GEMINI_TTS_VOICE"]
                    )
                )
            ),
        ),
    )
    pcm = next(
        (p.inline_data.data for p in (parts or [])
         if getattr(p, "inline_data", None) and p.inline_data.data),
        None,
    )
    if pcm is None:
        return {"ok": False, "error": "TTS não retornou áudio"}
    wav_bytes, duration = _pcm_to_wav(pcm)

    media_url = storage.save_bytes(user_id, wav_bytes, "wav")
    msg_id = _new_assistant_media(
        user_id, "", media_url, "audio",
        {"mime": "audio/wav", "duration": duration, "transcript": text},
    )
    return {"ok": True, "message_id": msg_id, "created": "áudio", "duration": duration}


def generate_document(user_id: int, args: dict) -> dict:
    fmt = (args.get("format") or "pdf").lower()
    title = (args.get("title") or "Documento").strip()
    body = args.get("content") or ""
    if fmt not in ("pdf", "docx", "xlsx", "txt"):
        return {"ok": False, "error": f"formato não suportado: {fmt}"}

    try:
        data = _render_document(fmt, title, body, args)
    except Exception as exc:  # noqa: BLE001
        return {"ok": False, "error": f"falha ao gerar {fmt}: {exc}"}

    media_type, mime = storage.classify(fmt)
    media_url = storage.save_bytes(user_id, data, fmt)
    msg_id = _new_assistant_media(
        user_id, title, media_url, media_type,
        {"mime": mime, "original_name": f"{title}.{fmt}"},
    )
    return {"ok": True, "message_id": msg_id, "created": f"documento {fmt}"}


def _render_document(fmt: str, title: str, body: str, args: dict) -> bytes:
    buf = io.BytesIO()
    if fmt == "txt":
        buf.write(f"{title}\n\n{body}".encode("utf-8"))
    elif fmt == "pdf":
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        doc = SimpleDocTemplate(buf, pagesize=A4)
        styles = getSampleStyleSheet()
        flow = [Paragraph(title, styles["Title"]), Spacer(1, 12)]
        for para in body.split("\n"):
            flow.append(Paragraph(para or "&nbsp;", styles["BodyText"]))
        doc.build(flow)
    elif fmt == "docx":
        from docx import Document

        d = Document()
        d.add_heading(title, level=1)
        for para in body.split("\n"):
            d.add_paragraph(para)
        d.save(buf)
    elif fmt == "xlsx":
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = title[:31] or "Planilha"
        # `rows`: lista de listas; senão, uma célula por linha do content
        rows = args.get("rows")
        if isinstance(rows, list):
            for row in rows:
                ws.append(row if isinstance(row, list) else [row])
        else:
            for line in body.split("\n"):
                ws.append([line])
        wb.save(buf)
    return buf.getvalue()
